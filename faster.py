import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches
import cv2
import numpy as np

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_blocks(page):

    text_blocks = page.get_text("blocks")
    return text_blocks

def extract_images(page, output_folder, page_number):
    images = page.get_images(full=True)
    image_paths = []
    for img_index, img in enumerate(images):
        xref = img[0]
        base_image = page.parent.extract_image(xref)
        image_data = base_image["image"]
        img_path = os.path.join(output_folder, f"page_{page_number}_image_{img_index}.png")
        with open(img_path, "wb") as img_file:
            img_file.write(image_data)
        image_paths.append(img_path)
    return image_paths

def detect_tables(image_path):
    img = cv2.imread(image_path, 0)
    _, binary = cv2.threshold(img, 200, 255, cv2.THRESH_BINARY)
    contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    table_regions = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if w > 50 and h > 50:
            table_regions.append((x, y, w, h))
    return table_regions

def generate_word_document(pdf_path, output_docx, temp_folder="temp"):
    doc = Document()
    os.makedirs(temp_folder, exist_ok=True)
    
    pdf_doc = fitz.open(pdf_path)
    for page_number in range(len(pdf_doc)):
        page = pdf_doc[page_number]
        doc.add_heading(f"Page {page_number + 1}", level=1)
        text_blocks = extract_text_blocks(page)
        for block in text_blocks:
            x0, y0, x1, y1, text, _ = block
            if text.strip():
                doc.add_paragraph(text.strip())
        image_paths = extract_images(page, temp_folder, page_number)
        for img_path in image_paths:
            doc.add_picture(img_path, width=Inches(5.0))

        pix = page.get_pixmap(dpi=300)
        img_path = os.path.join(temp_folder, f"page_{page_number}.png")
        pix.save(img_path)

        table_regions = detect_tables(img_path)
        for idx, (x, y, w, h) in enumerate(table_regions):
            table_img = Image.open(img_path).crop((x, y, x + w, y + h))
            table_text = pytesseract.image_to_string(table_img, config='--psm 6')
            doc.add_paragraph(f"Table {idx + 1}:\n{table_text}")


    doc.save(output_docx)
    print(f"Word document saved at {output_docx}")


pdf_path = "docs\Original.pdf"  
output_docx = "output.docx"
generate_word_document(pdf_path, output_docx)
