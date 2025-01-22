import os
from pdf2image import convert_from_path
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION

def flatten_pdf(input_pdf, output_pdf):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(output_pdf, "wb") as f:
        writer.write(f)
    print(f"Flattened PDF saved at: {output_pdf}")

def convert_pdf_to_docx(flatten_pdf, output_docx):
    cv = Converter(flatten_pdf)
    cv.convert(output_docx, start=0, end=None, layout_analysis=True, multi_processing=True)
    cv.close()
    print(f"Converted PDF to DOCX: {output_docx}")

def adjust_docx_formatting(docx_file, pdf_file):
    reader = PdfReader(pdf_file)
    pdf_page = reader.pages[0]
    pdf_width = pdf_page.mediabox.width / 72
    pdf_height = pdf_page.mediabox.height / 72

    doc = Document(docx_file)
    for section in doc.sections:
        section.page_width = Inches(pdf_width)
        section.page_height = Inches(pdf_height)
        section.orientation = (
            WD_ORIENTATION.LANDSCAPE if pdf_width > pdf_height else WD_ORIENTATION.PORTRAIT
        )
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.footer_distance = Inches(0.3)

    for table in doc.tables:
        for row in table.rows:
            row.allow_break_across_pages = False

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(6)

    doc.save(docx_file)
    print(f"Adjusted formatting for: {docx_file}")

def clean_up_folder(folder):
    if os.path.exists(folder):
        for file_name in os.listdir(folder):
            os.remove(os.path.join(folder, file_name))
        os.rmdir(folder)
        print(f"Removed folder: {folder}")

def process_pdfs(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)

    for file_name in os.listdir(input_folder):
        if file_name.endswith(".pdf"):
            input_pdf = os.path.join(input_folder, file_name)
            output_subfolder = os.path.join(output_folder, os.path.splitext(file_name)[0])
            os.makedirs(output_subfolder, exist_ok=True)

            flattened_pdf = os.path.join(output_subfolder, "flattened.pdf")
            docx_output = os.path.join(output_folder, f"{os.path.splitext(file_name)[0]}.docx")

            if os.path.exists(docx_output):
                print(f"\n=== Skipping {file_name}: Already Converted to DOCX ===")
                clean_up_folder(output_subfolder)
                continue

            print(f"\n=== Processing {file_name} ===")
            print("\n=== Flattening PDF ===")
            flatten_pdf(input_pdf, flattened_pdf)

            print("\n=== Converting PDF to DOCX ===")
            convert_pdf_to_docx(flattened_pdf, docx_output)

            print("\n=== Adjusting DOCX Formatting ===")
            adjust_docx_formatting(docx_output, flattened_pdf)

            print("\n=== Cleaning Up Intermediate Files ===")
            clean_up_folder(output_subfolder)

            print(f"\n=== Completed Processing for: {file_name} ===")

if __name__ == "__main__":
    process_pdfs("docs", "processed_output")
