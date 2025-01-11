import pdf2docx
import camelot
from docx import Document
from docx.shared import Inches

def convert_pdf_to_docx(pdf_path, docx_path):
    pdf2docx.parse(pdf_path, docx_path)

def extract_tables(pdf_path):
    tables = camelot.read_pdf(pdf_path, flavor='stream')
    return tables

def create_word_document(docx_path):
    doc = Document()
    return doc

def insert_tables_into_doc(doc, tables):
    for table in tables:
        # Add a table to the document
        doc_table = doc.add_table(rows=table.df.shape[0], cols=table.df.shape[1])
        for i in range(table.df.shape[0]):
            for j in range(table.df.shape[1]):
                doc_table.cell(i, j).text = str(table.df.iloc[i, j])
        doc.add_paragraph()  # Add space between tables

def full_conversion(pdf_path, docx_path):
    convert_pdf_to_docx(pdf_path, docx_path)
    tables = extract_tables(pdf_path)
    doc = Document(docx_path)
    insert_tables_into_doc(doc, tables)
    doc.save(docx_path)

if __name__ == "__main__":
    pdf_path = 'docs/Original.pdf'
    docx_path = 'output.docx'
    full_conversion(pdf_path, docx_path)