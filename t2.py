import os
from pdf2image import convert_from_path
from PyPDF2 import PdfReader, PdfWriter
import pytesseract
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches
from PIL import Image

def flatten_pdf(input_pdf, output_pdf):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(output_pdf, "wb") as f:
        writer.write(f)
    print(f"Flattened PDF saved at: {output_pdf}")

def rasterize_pdf(input_pdf, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    images = convert_from_path(input_pdf)
    for i, image in enumerate(images):
        image_path = os.path.join(output_folder, f"page_{i + 1}.png")
        image.save(image_path, "PNG")
        print(f"Saved rasterized page: {image_path}")
    return images

def ocr_pdf_images(image_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    for image_file in sorted(os.listdir(image_folder)):
        if image_file.endswith('.png'):
            image_path = os.path.join(image_folder, image_file)
            text = pytesseract.image_to_string(image_path, lang="eng+rus")
            output_file = os.path.join(output_folder, f"{image_file.replace('.png', '.txt')}")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"OCR text saved to: {output_file}")

def convert_pdf_to_docx(input_pdf, output_docx):
    cv = Converter(input_pdf)
    cv.convert(output_docx, start=0, end=None, layout_analysis=True)
    cv.close()
    print(f"Converted PDF to DOCX: {output_docx}")

def adjust_docx_formatting(docx_file):
    adjust_docx_margins(docx_file)
    fix_table_split(docx_file)

def adjust_docx_margins(docx_file):
    doc = Document(docx_file)
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    doc.save(docx_file)



def clean_up_intermediate_files(folders):
    for folder in folders:
        if os.path.exists(folder):
            for file_name in os.listdir(folder):
                file_path = os.path.join(folder, file_name)
                os.remove(file_path)
            os.rmdir(folder)
            print(f"Removed folder: {folder}")
            
def fix_table_split(docx_file):
    doc = Document(docx_file)
    for table in doc.tables:
        for row in table.rows:
            row.allow_break_across_pages = False
    doc.save(docx_file)

def process_pdfs_in_folder(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    for file_name in os.listdir(input_folder):
        if file_name.endswith(".pdf"):
            input_pdf = os.path.join(input_folder, file_name)
            output_subfolder = os.path.join(output_folder, os.path.splitext(file_name)[0])
            os.makedirs(output_subfolder, exist_ok=True)

            flattened_pdf = os.path.join(output_subfolder, "flattened.pdf")
            rasterized_folder = os.path.join(output_subfolder, "rasterized_images")
            ocr_output_folder = os.path.join(output_subfolder, "ocr_texts")
            docx_output = os.path.join(output_subfolder, f"{os.path.splitext(file_name)[0]}.docx")

            print(f"\n=== Processing {file_name} ===")
            flatten_pdf(input_pdf, flattened_pdf)
            rasterize_pdf(flattened_pdf, rasterized_folder)
            ocr_pdf_images(rasterized_folder, ocr_output_folder)
            convert_pdf_to_docx(flattened_pdf, docx_output)
            adjust_docx_formatting(docx_output)
            print(f"\n=== Completed Processing for: {file_name} ===")

            print("\n=== Cleaning Up Intermediate Files ===")
            if os.path.exists(flattened_pdf):
                os.remove(flattened_pdf)
                print(f"Deleted: {flattened_pdf}")

if __name__ == "__main__":
    input_folder = "docs1"  
    output_folder = "processed_output"  
    process_pdfs_in_folder(input_folder, output_folder)
