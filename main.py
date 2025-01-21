import os
from pdf2image import convert_from_path
from PyPDF2 import PdfReader, PdfWriter
import pytesseract
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches
from PIL import Image

def flatten_pdf(input_pdf, flatten_pdf):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(flatten_pdf, "wb") as f:
        writer.write(f)
    print(f"Flattened PDF saved at: {flatten_pdf}")

def convert_pdf_to_docx(flatten_pdf, output_docx):
    cv = Converter(flatten_pdf)
    cv.convert(output_docx, start=0, end=None, layout_analysis=True, multi_processing=True)
    cv.close()
    print(f"Converted PDF to DOCX: {output_docx}")

def adjust_docx_formatting(docx_file):
    adjust_docx_margins(docx_file)
    fix_table_split(docx_file)

def adjust_docx_margins(docx_file):
    doc = Document(docx_file)
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    doc.save(docx_file)
    print(f"Margins adjusted for: {docx_file}")

def fix_table_split(docx_file):
    doc = Document(docx_file)
    for table in doc.tables:
        for row in table.rows:
            row.allow_break_across_pages = True
    doc.save(docx_file)

def clean_up_folder(folder):
    if os.path.exists(folder):
        for file_name in os.listdir(folder):
            file_path = os.path.join(folder, file_name)
            os.remove(file_path)
        os.rmdir(folder)
        print(f"Removed folder: {folder}")

def process_pdfs_in_folder(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)

    for file_name in os.listdir(input_folder):
        if file_name.endswith(".pdf"):
            input_pdf = os.path.join(input_folder, file_name)
            output_subfolder = os.path.join(output_folder, os.path.splitext(file_name)[0])
            os.makedirs(output_subfolder, exist_ok=True)

            flattened_pdf = os.path.join(output_subfolder, "flattened.pdf")
            docx_output = os.path.join(output_folder, f"{os.path.splitext(file_name)[0]}.docx")

            if os.path.exists(docx_output):
                print(f"\n=== Skipping {file_name}: Already Converted to DOCX ===")
                print("\n=== Cleaning Up Intermediate Files ===")
                clean_up_folder(output_subfolder)                
                continue

            print(f"\n=== Processing {file_name} ===")

            print("\n=== Flattening PDF ===")
            flatten_pdf(input_pdf, flattened_pdf)

            print("\n=== Converting PDF to DOCX ===")
            convert_pdf_to_docx(flattened_pdf, docx_output)

            print("\n=== Adjusting DOCX Formatting ===")
            adjust_docx_formatting(docx_output)

            print("\n=== Cleaning Up Intermediate Files ===")
            clean_up_folder(output_subfolder)

            print(f"\n=== Completed Processing for: {file_name} ===")

if __name__ == "__main__":
    input_folder = "docs"  
    output_folder = "processed_output"  
    process_pdfs_in_folder(input_folder, output_folder)
