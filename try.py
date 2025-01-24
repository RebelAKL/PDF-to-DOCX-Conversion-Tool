import os
from pdf2image import convert_from_path
from PyPDF2 import PdfReader, PdfWriter
import pytesseract
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
import cProfile

def flatten_pdf(input_pdf, flatten_pdf):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(flatten_pdf, "wb") as f:
        writer.write(f)
    print(f"Flattened PDF saved at: {flatten_pdf}")

def convert_pdf_to_docx(flatten_pdf, output_docx):
    cv = Converter(flatten_pdf)
    cv.convert(output_docx, start=0, end=None, layout_analysis=True, multi_processing=True)
    cv.close()
    print(f"Converted PDF to DOCX: {output_docx}")

def adjust_docx_formatting(docx_file):
    doc = Document(docx_file)

    # Adjust margins
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)  
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(6) 

    doc.save(docx_file)
    print(f"Adjusted formatting (margins and font sizes) for: {docx_file}")

def clean_up_folder(folder):
    if os.path.exists(folder):
        for file_name in os.listdir(folder):
            os.remove(os.path.join(folder, file_name))
        os.rmdir(folder)
        print(f"Removed folder: {folder}")

def process_single_pdf(file_name, input_folder, output_folder):
    input_pdf = os.path.join(input_folder, file_name)
    output_subfolder = os.path.join(output_folder, os.path.splitext(file_name)[0])
    os.makedirs(output_subfolder, exist_ok=True)

    flattened_pdf = os.path.join(output_subfolder, "flattened.pdf")
    docx_output = os.path.join(output_folder, f"{os.path.splitext(file_name)[0]}.docx")

    if os.path.exists(docx_output):
        print(f"\n=== Skipping {file_name}: Already Converted to DOCX ===")
        print("\n=== Cleaning Up Intermediate Files ===")
        clean_up_folder(output_subfolder)
        return

    print(f"\n=== Processing {file_name} ===")
    print("\n=== Flattening PDF ===")
    flatten_pdf(input_pdf, flattened_pdf)

    print("\n=== Converting PDF to DOCX ===")
    convert_pdf_to_docx(flattened_pdf, docx_output)

    print("\n=== Adjusting DOCX Formatting ===")
    adjust_docx_formatting(docx_output)

    print("\n=== Cleaning Up Intermediate Files ===")
    clean_up_folder(output_subfolder)

    print(f"\n=== Completed Processing for: {file_name} ===")

def process_pdfs_in_folder(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    pdf_files = [file_name for file_name in os.listdir(input_folder) if file_name.endswith(".pdf")]

    print(f"\n=== Starting Parallel Processing for {len(pdf_files)} PDFs ===")
    with ThreadPoolExecutor() as executor:
        executor.map(lambda file_name: process_single_pdf(file_name, input_folder, output_folder), pdf_files)

if __name__ == "__main__":
    input_folder = "docs" 
    output_folder = "processed_output"
    
    # Profile the performance of the process
    print("\n=== Profiling Code Performance ===")
    cProfile.run("process_pdfs_in_folder(input_folder, output_folder)")